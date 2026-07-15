import io
from docx import Document

from app.schemas import ResumeData

def export_to_docx(resume_data: ResumeData) -> bytes:
    """Export resume data to DOCX format.
    
    Args:
        resume_data (ResumeData): The resume data to export.
        
    Returns:
        bytes: The DOCX file content as bytes.
    """
    doc = Document()
    
    # Add contact information
    doc.add_heading(f"{resume_data.contact.name}", level=1)
    doc.add_paragraph(f"Email: {resume_data.contact.email}")
    if resume_data.contact.phone:
        doc.add_paragraph(f"Phone: {resume_data.contact.phone}")
    if resume_data.contact.location:
        doc.add_paragraph(f"Location: {resume_data.contact.location}")
    
    # Add summary
    if resume_data.summary:
        doc.add_heading("Summary", level=2)
        doc.add_paragraph(resume_data.summary)
    
    # Add skills
    if resume_data.skills:
        doc.add_heading("Skills", level=2)
        for skill in resume_data.skills:
            doc.add_paragraph(skill)
    
    # Add experience
    if resume_data.experience:
        doc.add_heading("Work Experience", level=1)
        for exp in resume_data.experience:
            doc.add_heading(exp.title, level=2)
            doc.add_paragraph(f"{exp.company} ({exp.start_date} - {exp.end_date})")
            if exp.location:
                doc.add_paragraph(exp.location)
            for bullet in exp.bullets:
                doc.add_paragraph(bullet, style='List Bullet')
    
    # Add education
    if resume_data.education:
        doc.add_heading("Education", level=1)
        for edu in resume_data.education:
            doc.add_heading(f"{edu.degree} from {edu.institution}", level=2)
            if edu.year:
                doc.add_paragraph(f"Year: {edu.year}")
    
    # Add certifications
    if resume_data.certifications:
        doc.add_heading("Certifications", level=1)
        for cert in resume_data.certifications:
            doc.add_paragraph(cert)
    
    # Save the document to a bytes stream and return it
    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    return output.getvalue()