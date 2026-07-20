"""Export a structured resume to Markdown, DOCX and PDF.

DOCX and PDF use the same targeted-resume design: Arial-compatible fonts,
compact spacing, tab-aligned dates, bottom-ruled headings, and US Letter pages.
"""
import logging
import re
from pathlib import Path

import fitz  # PyMuPDF
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Mm

from dataclasses import dataclass

from app.schemas import ResumeData

logger = logging.getLogger(__name__)

NAVY = "185076"
GREY = "68767D"
DARK = "1C1C1C"
FONT_NAME = "Arial"


@dataclass
class ExportTheme:
    """Configurable visual theme for PDF/DOCX export."""

    name: str = "Classic"
    font_family: str = "Arial"
    accent_color: str = "185076"
    secondary_color: str = "68767D"
    text_color: str = "1C1C1C"
    body_font_size: float = 9.0
    heading_font_size: float = 10.5
    title_font_size: float = 20.5
    line_height_body: float = 10.0
    line_height_compact: float = 9.5
    margin_top_bottom: float = 36.0
    margin_left_right: float = 43.2

    @property
    def accent_rgb(self) -> RGBColor:
        h = self.accent_color
        return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))

    @property
    def secondary_rgb(self) -> RGBColor:
        h = self.secondary_color
        return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))

    @property
    def text_rgb(self) -> RGBColor:
        h = self.text_color
        return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))

    def accent_float(self) -> tuple[float, float, float]:
        h = self.accent_color
        return (int(h[0:2], 16) / 255, int(h[2:4], 16) / 255, int(h[4:6], 16) / 255)

    def secondary_float(self) -> tuple[float, float, float]:
        h = self.secondary_color
        return (int(h[0:2], 16) / 255, int(h[2:4], 16) / 255, int(h[4:6], 16) / 255)

    def text_float(self) -> tuple[float, float, float]:
        h = self.text_color
        return (int(h[0:2], 16) / 255, int(h[2:4], 16) / 255, int(h[4:6], 16) / 255)


CLASSIC_THEME = ExportTheme(name="Classic")
MODERN_THEME = ExportTheme(
    name="Modern",
    font_family="Calibri",
    accent_color="1A3C5E",
    secondary_color="5A6A72",
    text_color="1C1C1C",
    body_font_size=9.0,
    heading_font_size=10.5,
    title_font_size=20.5,
)
COMPACT_THEME = ExportTheme(
    name="Compact",
    body_font_size=9.5,
    heading_font_size=10.5,
    title_font_size=20.0,
    line_height_body=9.5,
    line_height_compact=9.0,
    margin_top_bottom=30.0,
    margin_left_right=36.0,
)

TEMPLATES: dict[str, ExportTheme] = {
    "Classic": CLASSIC_THEME,
    "Modern": MODERN_THEME,
    "Compact": COMPACT_THEME,
}


def get_template(name: str) -> ExportTheme:
    """Return the template by name, falling back to Classic."""
    return TEMPLATES.get(name, CLASSIC_THEME)


# Legacy color constants for backward compatibility
_NAVY_RGB = CLASSIC_THEME.accent_rgb
_GREY_RGB = CLASSIC_THEME.secondary_rgb
_DARK_RGB = CLASSIC_THEME.text_rgb

_NAVY_F = CLASSIC_THEME.accent_float()
_GREY_F = CLASSIC_THEME.secondary_float()
_DARK_F = CLASSIC_THEME.text_float()

SEP = "  |  "
EXP_SEP = " | "

# US Letter geometry used by the targeted resume template.
PAGE_WIDTH_MM = 215.9
PAGE_HEIGHT_MM = 279.4
PAGE_WIDTH_PT = 612.0
PAGE_HEIGHT_PT = 792.0
MARGIN_TOP_BOTTOM_PT = 36.0
MARGIN_LEFT_RIGHT_PT = 43.2
ROLE_DATE_TAB_PT = 507.6
CERT_ISSUER_TAB_PT = 392.4
CERT_YEAR_TAB_PT = 507.6


def _date_range(start_date: str = "", end_date: str = "") -> str:
    """Return a consistently formatted date range."""
    return " \u2013 ".join(filter(None, [start_date.strip(), end_date.strip()]))


def _docx_date_text(value: str) -> str:
    """Keep a DOCX date range together when a heading is close to full width."""
    return value.replace(" ", "\u00a0")


def _normalise_date_text(value: str) -> str:
    """Normalise separators so equivalent date ranges compare equal."""
    return re.sub(r"\s*[-\u2013\u2014\u00b7]\s*", " - ", value.strip()).casefold()


_DATE_RANGE_TEXT_RE = re.compile(
    r"\b(?:19|20)\d{2}\b.*(?:\b(?:19|20)\d{2}\b|\bpresent\b|\bcurrent\b)",
    re.I,
)
_PROJECT_CONTEXT_RE = re.compile(
    r"^(.*?)\s+[\u00b7|]\s+(Project\s+(?:within|under|for)\b.*)$",
    re.I,
)
SUMMARY_MAX_WORDS = 300
_EMBEDDED_SKILLS_RE = re.compile(
    r"\s+(?:CORE\s+)?TECHNICAL\s+SKILLS\s*:?[ \t]+"
    r"(?=(?:Routing\s*&\s*Switching|Security|SD-WAN|Wireless|"
    r"Project\s+Delivery|Diagnostics)\b)",
    re.I,
)
_SKILL_CATEGORY_RE = re.compile(
    r"(?P<label>Routing\s*&\s*Switching|Security|SD-WAN|Wireless|"
    r"Project\s+Delivery\s*&\s*Documentation|Diagnostics\s*&\s*Infrastructure)"
    r"\s*:\s*",
    re.I,
)


def _truncate_summary(text: str, max_words: int = SUMMARY_MAX_WORDS) -> str:
    """Return a clean summary no longer than *max_words*.

    Prefer a complete sentence near the limit; otherwise close the truncated
    text with a period so the exported resume never ends mid-thought.
    """
    clean = re.sub(r"\s+", " ", text or "").strip()
    words = clean.split()
    if len(words) <= max_words:
        return clean

    candidate = " ".join(words[:max_words])
    sentence_ends = [candidate.rfind(mark) for mark in (". ", "! ", "? ")]
    last_end = max(sentence_ends)
    if last_end >= len(candidate) * 0.6:
        return candidate[: last_end + 1].strip()
    return candidate.rstrip(" ,;:-") + "."


def _parse_embedded_skills(text: str) -> list[str]:
    """Split a flattened categorized skills block into exportable entries."""
    clean = re.sub(r"\s+", " ", text or "").strip()
    if not clean:
        return []

    matches = list(_SKILL_CATEGORY_RE.finditer(clean))
    if not matches:
        return [clean]

    entries: list[str] = []
    for index, match in enumerate(matches):
        start = match.end()
        end = matches[index + 1].start() if index + 1 < len(matches) else len(clean)
        value = clean[start:end].strip(" \t,;•")
        label = re.sub(r"\s+", " ", match.group("label")).strip()
        if value:
            entries.append(f"{label}: {value}")
    return entries


def _summary_and_skills(resume: ResumeData) -> tuple[str, list[str]]:
    """Keep accidentally flattened skills out of the summary.

    Some AI/parser outputs append an uppercase ``CORE TECHNICAL SKILLS`` block
    to ``summary`` and leave ``skills`` empty. Recover that block at export
    time, while preserving a populated canonical skills list when available.
    """
    raw_summary = resume.summary or ""
    embedded_skills = ""
    match = _EMBEDDED_SKILLS_RE.search(raw_summary)
    if match:
        embedded_skills = raw_summary[match.end():]
        raw_summary = raw_summary[:match.start()]

    summary = _truncate_summary(raw_summary)
    skills = [re.sub(r"\s+", " ", skill).strip() for skill in resume.skills if skill.strip()]
    if not skills and embedded_skills:
        skills = _parse_embedded_skills(embedded_skills)
    return summary, skills


def _skill_label_value(entry: str) -> tuple[str, str] | None:
    """Return ``(label, value)`` for a categorized skill entry."""
    if ":" not in entry:
        return None
    label, value = entry.split(":", 1)
    label, value = label.strip(), value.strip()
    if not label or not value or len(label) > 50:
        return None
    return label, value


def _project_display_parts(project) -> tuple[str, str, str]:
    """Return project title, contextual metadata, and date range.

    Older parsed resumes stored the date range in ``meta``. Newer records may
    have dedicated ``start_date`` / ``end_date`` fields while keeping ``meta``
    for contextual text such as "Project within Acme role". Support both.
    """
    explicit_dates = _date_range(project.start_date, project.end_date)
    meta = (project.meta or "").strip()
    title = project.title
    context = ""

    if explicit_dates:
        if meta and _normalise_date_text(meta) != _normalise_date_text(explicit_dates):
            context = meta
        dates = explicit_dates
    elif meta and _DATE_RANGE_TEXT_RE.search(meta):
        dates = meta
    else:
        context = meta
        dates = ""

    # Some older optimizer outputs embedded the context directly in title
    # while using ``meta`` for the dates. Pull that context onto its own compact
    # line so a long heading cannot force the date outside the right margin.
    if not context:
        match = _PROJECT_CONTEXT_RE.match(title)
        if match:
            title, context = match.group(1).strip(), match.group(2).strip()

    return title, context, dates


def _bottom_rule(paragraph, color: str = NAVY) -> None:
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), color)
    pBdr.append(bottom)
    pPr.append(pBdr)


def _blocks(resume: ResumeData) -> list[tuple[str, str]]:
    """Return Markdown blocks in the same order as the targeted template."""
    summary, skills = _summary_and_skills(resume)
    blocks: list[tuple[str, str]] = [("h1", resume.contact.name or "Resume")]
    if resume.headline:
        blocks.append(("meta", resume.headline))
    contact = SEP.join(filter(None, [
        resume.contact.location,
        resume.contact.phone,
        resume.contact.email,
        resume.contact.linkedin,
        resume.contact.website,
    ]))
    if contact:
        blocks.append(("meta", contact))
    if summary:
        blocks += [("h2", "Professional Summary"), ("p", summary)]
    if skills:
        blocks += [("h2", "Core Technical Skills"), ("p", ", ".join(skills))]
    if resume.experience:
        blocks.append(("h2", "Professional Experience"))
        for exp in resume.experience:
            metadata = " | ".join(filter(None, [exp.company, exp.location]))
            dates = _date_range(exp.start_date, exp.end_date)
            header = exp.title + (f" | {metadata}" if metadata else "")
            if dates:
                header += f" | {dates}"
            blocks.append(("h3", header))
            for bullet in exp.bullets:
                blocks.append(("li", bullet))
    if resume.projects:
        blocks.append(("h2", "Selected Project Delivery"))
        for project in resume.projects:
            title, context, dates = _project_display_parts(project)
            header = title + (f" | {context}" if context else "")
            if dates:
                header += f" | {dates}"
            blocks.append(("h3", header))
            for bullet in project.bullets:
                blocks.append(("li", bullet))
    if resume.certifications:
        blocks.append(("h2", "Certifications"))
        for certification in resume.certifications:
            title, issuer, year = _certification_parts(certification)
            blocks.append(("p", " | ".join(filter(None, [title, issuer, year]))))
    if resume.education:
        blocks.append(("h2", "Education"))
        for education in resume.education:
            degree, detail = _education_display(education)
            blocks.append(("p", " ".join(filter(None, [degree, detail]))))
    if resume.languages:
        blocks += [("h2", "Languages"), ("p", " | ".join(resume.languages))]
    return blocks


def to_markdown(resume: ResumeData) -> str:
    out: list[str] = []
    for kind, text in _blocks(resume):
        if kind == "h1":
            out.append(f"# {text}")
        elif kind == "h2":
            out.append(f"\n## {text}\n")
        elif kind == "h3":
            out.append(f"### {text}")
        elif kind == "li":
            out.append(f"- {text}")
        else:
            out.append(text)
    return "\n".join(out).strip() + "\n"


def export_markdown(resume: ResumeData, path: str) -> None:
    logger.info("Exporting resume to Markdown: %s", path)
    Path(path).write_text(to_markdown(resume), encoding="utf-8")


def _set_style_font(style, *, size: float | None = None, color=None, bold=None, font_name: str | None = None) -> None:
    """Apply the given font consistently to a DOCX style, including East Asian slots."""
    style.font.name = font_name or FONT_NAME
    if size is not None:
        style.font.size = Pt(size)
    if color is not None:
        style.font.color.rgb = color
    if bold is not None:
        style.font.bold = bold
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    actual = font_name or FONT_NAME
    for attr in ("ascii", "hAnsi", "cs", "eastAsia"):
        rFonts.set(qn(f"w:{attr}"), actual)


def _compact_year_range(value: str) -> str:
    """Use an en dash without surrounding spaces in compact education rows."""
    return re.sub(r"\s*[-\u2013\u2014]\s*", "\u2013", (value or "").strip())


def _education_display(education) -> tuple[str, str]:
    """Return the bold degree label and the remaining reference-template text."""
    degree = (education.degree or "").strip().rstrip(":")
    institution = re.sub(r"\s*,\s*(GPA\s*[:]?\s*)", r" | \1", education.institution or "", flags=re.I)
    institution = re.sub(r"GPA\s*:\s*", "GPA ", institution, flags=re.I).strip()
    year = _compact_year_range(education.year)
    tail = institution
    if year:
        tail = f"{tail} | {year}" if tail else year
    return (degree + ":" if degree else ""), tail


def _certification_parts(certification: str) -> tuple[str, str, str]:
    """Split a certification into title, issuer, and year columns."""
    parts = [part.strip() for part in certification.split("|") if part.strip()]
    if len(parts) >= 3:
        return parts[0], parts[1], parts[2]
    if len(parts) == 2:
        year_match = re.fullmatch(r"(?:19|20)\d{2}", parts[1])
        return (parts[0], "", parts[1]) if year_match else (parts[0], parts[1], "")
    text = parts[0] if parts else certification.strip()
    year_match = re.search(r"\b((?:19|20)\d{2})\s*$", text)
    if year_match:
        return text[:year_match.start()].rstrip(" -\u00b7|\t"), "", year_match.group(1)
    return text, "", ""


def export_docx(resume: ResumeData, path: str, theme: ExportTheme | None = None) -> None:
    """Export using the compact targeted resume template."""
    t = theme or CLASSIC_THEME
    logger.info("Exporting resume to DOCX [%s]: %s", t.name, path)
    doc = Document()

    accent_hex = t.accent_color
    accent_rgb = t.accent_rgb
    secondary_rgb = t.secondary_rgb
    text_rgb = t.text_rgb

    section = doc.sections[0]
    section.page_width = Mm(PAGE_WIDTH_MM)
    section.page_height = Mm(PAGE_HEIGHT_MM)
    section.top_margin = Pt(t.margin_top_bottom)
    section.bottom_margin = Pt(t.margin_top_bottom)
    section.left_margin = Pt(t.margin_left_right)
    section.right_margin = Pt(t.margin_left_right)

    normal = doc.styles["Normal"]
    _set_style_font(normal, size=t.body_font_size, color=text_rgb, font_name=t.font_family)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(1.1)
    normal.paragraph_format.line_spacing = 1.0

    compact = doc.styles.add_style("Compact", WD_STYLE_TYPE.PARAGRAPH)
    _set_style_font(compact, size=t.body_font_size, color=text_rgb, font_name=t.font_family)
    compact.paragraph_format.space_before = Pt(0)
    compact.paragraph_format.space_after = Pt(0.55)
    compact.paragraph_format.line_spacing = 1.0

    section_style = doc.styles.add_style("Resume Section", WD_STYLE_TYPE.PARAGRAPH)
    _set_style_font(section_style, size=t.heading_font_size, color=accent_rgb, bold=True, font_name=t.font_family)
    section_style.paragraph_format.space_before = Pt(5.2)
    section_style.paragraph_format.space_after = Pt(2)
    section_style.paragraph_format.keep_with_next = True

    role_style = doc.styles.add_style("Resume Role", WD_STYLE_TYPE.PARAGRAPH)
    _set_style_font(role_style, size=9.5, color=text_rgb, font_name=t.font_family)
    role_style.paragraph_format.space_before = Pt(2)
    role_style.paragraph_format.space_after = Pt(0.4)
    role_style.paragraph_format.keep_with_next = True

    bullet_style = doc.styles.add_style("Resume Bullet", WD_STYLE_TYPE.PARAGRAPH)
    _set_style_font(bullet_style, size=t.body_font_size, color=text_rgb, font_name=t.font_family)
    bullet_style.paragraph_format.space_before = Pt(0)
    bullet_style.paragraph_format.space_after = Pt(0.65)
    bullet_style.paragraph_format.line_spacing = 1.0
    bullet_style.paragraph_format.left_indent = Pt(13.7)
    bullet_style.paragraph_format.first_line_indent = Pt(-9.35)

    def add_run(paragraph, text, *, color=None, size=None, bold=False, italic=False):
        run = paragraph.add_run(text)
        run.font.name = t.font_family
        if size is not None:
            run.font.size = Pt(size)
        run.font.color.rgb = color if color is not None else text_rgb
        run.font.bold = bold
        run.font.italic = italic
        return run

    def section_heading(text: str, *, add_spacer: bool = True, page_break: bool = False):
        if add_spacer:
            spacer_p = doc.add_paragraph(style="Resume Section")
            _bottom_rule(spacer_p, accent_hex)
        paragraph = doc.add_paragraph(style="Resume Section")
        paragraph.paragraph_format.page_break_before = page_break
        add_run(paragraph, text.upper(), color=accent_rgb, size=t.heading_font_size, bold=True)
        _bottom_rule(paragraph, accent_hex)
        return paragraph

    def add_role_line(title: str, metadata: list[str], dates: str = ""):
        paragraph = doc.add_paragraph(style="Resume Role")
        paragraph.paragraph_format.tab_stops.add_tab_stop(Pt(ROLE_DATE_TAB_PT), WD_TAB_ALIGNMENT.RIGHT)
        add_run(paragraph, title, bold=True)
        for item in [item for item in metadata if item]:
            add_run(paragraph, " | ")
            add_run(paragraph, item, color=secondary_rgb)
        if dates:
            add_run(paragraph, "\t")
            add_run(paragraph, dates, color=secondary_rgb, italic=True)
        return paragraph

    def add_bullet(text: str):
        paragraph = doc.add_paragraph(style="Resume Bullet")
        add_run(paragraph, "\u2022  " + text)
        return paragraph

    summary, skills = _summary_and_skills(resume)

    name = doc.add_paragraph()
    name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name.paragraph_format.space_after = Pt(0)
    add_run(name, (resume.contact.name or "Resume").upper(), color=accent_rgb, size=t.title_font_size, bold=True)

    spacer = doc.add_paragraph()
    spacer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    spacer.paragraph_format.space_after = Pt(0.4)

    if resume.headline:
        headline = doc.add_paragraph()
        headline.alignment = WD_ALIGN_PARAGRAPH.CENTER
        headline.paragraph_format.space_after = Pt(0.4)
        add_run(headline, resume.headline, color=secondary_rgb, size=t.heading_font_size)

    contact = SEP.join(filter(None, [
        resume.contact.location,
        resume.contact.phone,
        resume.contact.email,
        resume.contact.linkedin,
        resume.contact.website,
    ]))
    if contact:
        contact_p = doc.add_paragraph()
        contact_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        contact_p.paragraph_format.space_after = Pt(3)
        add_run(contact_p, contact, color=secondary_rgb, size=8.5)

    if summary:
        section_heading("Professional Summary")
        paragraph = doc.add_paragraph(style="Compact")
        paragraph.paragraph_format.space_after = Pt(1.4)
        add_run(paragraph, summary, size=t.body_font_size)

    if skills:
        section_heading("Core Technical Skills")
        categorized = [_skill_label_value(skill) for skill in skills]
        if len(skills) > 1 and all(categorized):
            for label, value in categorized:
                paragraph = doc.add_paragraph(style="Compact")
                add_run(paragraph, label + ": ", bold=True)
                add_run(paragraph, value)
        else:
            paragraph = doc.add_paragraph(style="Compact")
            add_run(paragraph, " | ".join(skills))

    if resume.experience:
        section_heading("Professional Experience")
        for experience in resume.experience:
            add_role_line(
                experience.title,
                [experience.company, experience.location],
                _date_range(experience.start_date, experience.end_date),
            )
            for bullet in experience.bullets:
                add_bullet(bullet)

    if resume.projects:
        section_heading("Selected Project Delivery", add_spacer=False, page_break=True)
        for project in resume.projects:
            project_title, project_context, project_dates = _project_display_parts(project)
            add_role_line(project_title, [project_context], project_dates)
            for bullet in project.bullets:
                add_bullet(bullet)

    if resume.certifications:
        section_heading("Certifications")
        for certification in resume.certifications:
            title, issuer, year = _certification_parts(certification)
            paragraph = doc.add_paragraph(style="Compact")
            paragraph.paragraph_format.tab_stops.add_tab_stop(Pt(CERT_ISSUER_TAB_PT), WD_TAB_ALIGNMENT.LEFT)
            paragraph.paragraph_format.tab_stops.add_tab_stop(Pt(CERT_YEAR_TAB_PT), WD_TAB_ALIGNMENT.RIGHT)
            add_run(paragraph, title, bold=True)
            if issuer or year:
                add_run(paragraph, "\t")
                add_run(paragraph, issuer, color=_GREY_RGB)
                add_run(paragraph, "\t")
                add_run(paragraph, year, color=_GREY_RGB)

    if resume.education:
        section_heading("Education")
        for education in resume.education:
            degree, detail = _education_display(education)
            paragraph = doc.add_paragraph(style="Compact")
            if degree:
                add_run(paragraph, degree + " ", bold=True)
            add_run(paragraph, detail)

    if resume.languages:
        section_heading("Languages")
        paragraph = doc.add_paragraph(style="Compact")
        language_text = " | ".join(resume.languages)
        if ":" in language_text:
            label, value = language_text.split(":", 1)
            add_run(paragraph, label + ":", bold=True)
            add_run(paragraph, value)
        else:
            add_run(paragraph, language_text)

    doc.save(path)


_PDF_SIZES = {"h1": 26.0, "h2": 10.5, "h3": 10.5, "li": 9.5, "p": 9.5, "meta": 9.0}
_PDF_COLORS = {"h1": _NAVY_F, "h2": _NAVY_F, "h3": _DARK_F, "li": _DARK_F, "p": _DARK_F, "meta": _GREY_F}


def _resolve_pdf_font(name: str) -> str:
    """Return *name* if it is a valid built-in font, else fall back to helv."""
    try:
        fitz.get_text_length("x", fontname=name, fontsize=10)
        return name
    except Exception:
        logger.debug("PyMuPDF font %r unavailable, falling back to helv", name)
        return "helv"


# Use Arimo (an Arial-compatible metric substitute) when available so the
# PDF supports Unicode punctuation such as en dashes. Fall back to built-ins.
_PDF_FONT_FILES = {
    "arimo-regular": "/usr/share/fonts/truetype/croscore/Arimo-Regular.ttf",
    "arimo-bold": "/usr/share/fonts/truetype/croscore/Arimo-Bold.ttf",
    "arimo-italic": "/usr/share/fonts/truetype/croscore/Arimo-Italic.ttf",
}
_PDF_REGULAR = "arimo-regular" if Path(_PDF_FONT_FILES["arimo-regular"]).exists() else _resolve_pdf_font("helv")
_PDF_BOLD = "arimo-bold" if Path(_PDF_FONT_FILES["arimo-bold"]).exists() else _resolve_pdf_font("hebo")
_PDF_ITALIC = "arimo-italic" if Path(_PDF_FONT_FILES["arimo-italic"]).exists() else _resolve_pdf_font("heit")


def export_pdf(resume: ResumeData, path: str, theme: ExportTheme | None = None, target_pages: int = 1) -> None:
    """Export a PDF matching the compact targeted DOCX template.

    When *target_pages* is set, the exporter will reduce font sizes and
    regenerate the PDF until the output fits within that page count (minimum
    font size 7.5 pt).
    """
    t = theme or CLASSIC_THEME
    accent_f = t.accent_float()
    secondary_f = t.secondary_float()
    text_f = t.text_float()

    def _build_pdf(body_size: float) -> fitz.Document:
        """Internal: build a PDF with the given body font size."""
        doc = fitz.open()
        page = doc.new_page(width=PAGE_WIDTH_PT, height=PAGE_HEIGHT_PT)
        left = t.margin_left_right
        right = PAGE_WIDTH_PT - t.margin_left_right
        top = t.margin_top_bottom
        bottom = PAGE_HEIGHT_PT - t.margin_top_bottom
        y = top
        font_cache: dict[str, fitz.Font] = {}

        def text_width(text: str, fontname: str, fontsize: float) -> float:
            font = font_cache.get(fontname)
            if font is None:
                font_file = _PDF_FONT_FILES.get(fontname)
                font = fitz.Font(fontfile=font_file) if font_file else fitz.Font(fontname=fontname)
                font_cache[fontname] = font
            return font.text_length(text, fontsize=fontsize)

        def new_page() -> None:
            nonlocal page, y
            page = doc.new_page(width=PAGE_WIDTH_PT, height=PAGE_HEIGHT_PT)
            y = top

        def ensure(needed: float) -> None:
            if y + needed > bottom:
                new_page()

        def insert(x: float, baseline: float, text: str, *, size=body_size, color=text_f, bold=False, italic=False):
            font = _PDF_BOLD if bold else (_PDF_ITALIC if italic else _PDF_REGULAR)
            page.insert_text(
                (x, baseline),
                text,
                fontsize=size,
                fontname=font,
                fontfile=_PDF_FONT_FILES.get(font),
                color=color,
            )

        def centered(text: str, *, size: float, color, bold=False):
            nonlocal y
            font = _PDF_BOLD if bold else _PDF_REGULAR
            x = (PAGE_WIDTH_PT - text_width(text, font, size)) / 2
            insert(x, y + size, text, size=size, color=color, bold=bold)
            y += size + 1.5

        def section_title(text: str, *, compact_top: bool = False):
            nonlocal y
            ensure(22)
            y += 5.2 if compact_top else 18.0
            baseline = y + t.heading_font_size
            insert(left, baseline, text.upper(), size=t.heading_font_size, color=accent_f, bold=True)
            line_y = baseline + 2.2
            page.draw_line((left, line_y), (right, line_y), color=accent_f, width=0.8)
            y = line_y + 4.2

        def wrapped_tokens(tokens, max_width, size=body_size):
            lines = []
            line = []
            width = 0.0
            for text, bold, color, italic in tokens:
                words = re.findall(r"\s+|\S+", text)
                font = _PDF_BOLD if bold else (_PDF_ITALIC if italic else _PDF_REGULAR)
                for word in words:
                    word_width = text_width(word, font, size)
                    if line and width + word_width > max_width:
                        lines.append(line)
                        line = []
                        width = 0.0
                    line.append((word, bold, color, italic))
                    width += word_width
            if line:
                lines.append(line)
            return lines or [[]]

        def draw_wrapped(tokens, *, x=left, max_width=None, size=body_size, line_height=t.line_height_body, hanging=0.0):
            nonlocal y
            width = max_width if max_width is not None else right - x
            lines = wrapped_tokens(tokens, width, size)
            ensure(line_height * len(lines) + 1)
            for line_index, line in enumerate(lines):
                cursor = x if line_index == 0 else x + hanging
                baseline = y + size
                for text, bold, color, italic in line:
                    insert(cursor, baseline, text, size=size, color=color, bold=bold, italic=italic)
                    font = _PDF_BOLD if bold else (_PDF_ITALIC if italic else _PDF_REGULAR)
                    cursor += text_width(text, font, size)
                y += line_height
            return len(lines)

        def role_line(title: str, metadata: list[str], dates: str = ""):
            nonlocal y
            ensure(13)
            role_size = body_size + 0.5
            tokens = [(title, True, text_f, False)]
            for item in [item for item in metadata if item]:
                tokens.extend([(" | ", False, text_f, False), (item, False, secondary_f, False)])
            date_width = text_width(dates, _PDF_ITALIC, body_size) if dates else 0
            available = (right - left) - date_width - (10 if dates else 0)
            lines = wrapped_tokens(tokens, available, role_size)
            for index, line in enumerate(lines):
                baseline = y + role_size
                cursor = left
                for text, bold, color, italic in line:
                    insert(cursor, baseline, text, size=role_size, color=color, bold=bold, italic=italic)
                    font = _PDF_BOLD if bold else (_PDF_ITALIC if italic else _PDF_REGULAR)
                    cursor += text_width(text, font, role_size)
                if index == 0 and dates:
                    insert(right - date_width, baseline, dates, size=body_size, color=secondary_f, italic=True)
                y += 10.2
            y += 0.4

        def bullet(text: str):
            nonlocal y
            bullet_x = left + 4.35
            text_x = left + 13.7
            lines = wrapped_tokens([(text, False, text_f, False)], right - text_x, body_size)
            ensure(len(lines) * 9.65 + 1)
            for index, line in enumerate(lines):
                baseline = y + body_size
                if index == 0:
                    insert(bullet_x, baseline, "\u2022", size=body_size)
                cursor = text_x
                for fragment, bold, color, italic in line:
                    insert(cursor, baseline, fragment, size=body_size, color=color, bold=bold, italic=italic)
                    cursor += text_width(fragment, _PDF_REGULAR, body_size)
                y += 9.65

        summary, skills = _summary_and_skills(resume)

        centered((resume.contact.name or "Resume").upper(), size=t.title_font_size, color=accent_f, bold=True)
        y += 8.5
        if resume.headline:
            centered(resume.headline, size=t.heading_font_size, color=secondary_f)
        contact = SEP.join(filter(None, [
            resume.contact.location,
            resume.contact.phone,
            resume.contact.email,
            resume.contact.linkedin,
            resume.contact.website,
        ]))
        if contact:
            centered(contact, size=8.5, color=secondary_f)
        y += 3

        if summary:
            section_title("Professional Summary")
            draw_wrapped([(summary, False, text_f, False)], size=body_size, line_height=t.line_height_body)
            y += 1.4

        if skills:
            section_title("Core Technical Skills")
            categorized = [_skill_label_value(skill) for skill in skills]
            if len(skills) > 1 and all(categorized):
                for label, value in categorized:
                    draw_wrapped([
                        (label + ": ", True, text_f, False),
                        (value, False, text_f, False),
                    ], size=body_size, line_height=t.line_height_compact)
            else:
                draw_wrapped([(" | ".join(skills), False, text_f, False)], size=body_size, line_height=t.line_height_compact)

        if resume.experience:
            section_title("Professional Experience")
            for experience in resume.experience:
                role_line(
                    experience.title,
                    [experience.company, experience.location],
                    _date_range(experience.start_date, experience.end_date),
                )
                for item in experience.bullets:
                    bullet(item)

        if resume.projects:
            if y > top:
                new_page()
            section_title("Selected Project Delivery", compact_top=True)
            for project in resume.projects:
                title, context, dates = _project_display_parts(project)
                role_line(title, [context], dates)
                for item in project.bullets:
                    bullet(item)

        if resume.certifications:
            section_title("Certifications")
            for certification in resume.certifications:
                ensure(9.7)
                title, issuer, year = _certification_parts(certification)
                baseline = y + body_size
                insert(left, baseline, title, size=body_size, bold=True)
                if issuer:
                    insert(left + CERT_ISSUER_TAB_PT, baseline, issuer, size=body_size, color=secondary_f)
                if year:
                    year_width = text_width(year, _PDF_REGULAR, body_size)
                    insert(left + CERT_YEAR_TAB_PT - year_width, baseline, year, size=body_size, color=secondary_f)
                y += 9.5

        if resume.education:
            section_title("Education")
            for education in resume.education:
                degree, detail = _education_display(education)
                draw_wrapped([
                    ((degree + " ") if degree else "", True, text_f, False),
                    (detail, False, text_f, False),
                ], size=body_size, line_height=t.line_height_compact)

        if resume.languages:
            section_title("Languages")
            text = " | ".join(resume.languages)
            if ":" in text:
                label, value = text.split(":", 1)
                draw_wrapped([
                    (label + ":", True, text_f, False),
                    (value, False, text_f, False),
                ], size=body_size, line_height=t.line_height_compact)
            else:
                draw_wrapped([(text, False, text_f, False)], size=body_size, line_height=t.line_height_compact)

        return doc

    logger.info("Exporting resume to PDF [%s]: %s", t.name, path)
    body_size = t.body_font_size
    min_size = 7.5
    doc = None

    while body_size >= min_size:
        if doc is not None:
            doc.close()
        doc = _build_pdf(body_size)
        if len(doc) <= target_pages:
            break
        body_size -= 0.5
        logger.debug("PDF overflows target %d pages, reducing font to %.1f", target_pages, body_size)

    doc.save(path)
    doc.close()


def export_text_docx(text: str, path: str, theme: ExportTheme | None = None) -> None:
    """Save plain text (e.g. a cover letter) as a US Letter DOCX file."""
    t = theme or CLASSIC_THEME
    logger.info("Exporting text to DOCX [%s]: %s", t.name, path)
    doc = Document()
    section = doc.sections[0]
    section.page_width = Mm(PAGE_WIDTH_MM)
    section.page_height = Mm(PAGE_HEIGHT_MM)
    for paragraph in text.split("\n"):
        doc.add_paragraph(paragraph)
    doc.save(path)
