import fitz
from docx import Document

from app.exports.exporter import (
    CLASSIC_THEME,
    COMPACT_THEME,
    MODERN_THEME,
    ExportTheme,
    TEMPLATES,
    export_docx,
    export_pdf,
    get_template,
    to_markdown,
)
from app.schemas import (
    ContactInfo,
    EducationItem,
    ExperienceItem,
    ProjectItem,
    ResumeData,
)


def _resume() -> ResumeData:
    return ResumeData(
        contact=ContactInfo(name="Jane Doe", email="jane@example.com"),
        summary="Backend developer.",
        skills=["Python", "SQL"],
        experience=[
            ExperienceItem(
                title="Developer",
                company="Acme",
                start_date="2020",
                end_date="Present",
                bullets=["Built things"],
            )
        ],
        education=[EducationItem(degree="BSc CS", institution="MIT", year="2019")],
        certifications=["AWS SAA"],
    )


def test_markdown_structure():
    md = to_markdown(_resume())
    assert "# Jane Doe" in md
    assert "## Professional Summary" in md
    assert "## Core Technical Skills" in md
    assert "Python, SQL" in md
    assert "### Developer | Acme | 2020 – Present" in md
    assert "- Built things" in md
    assert "## Education" in md
    assert "AWS SAA" in md


def test_markdown_minimal_resume():
    md = to_markdown(ResumeData())
    assert md.startswith("# Resume")


def _layout_resume() -> ResumeData:
    return ResumeData(
        contact=ContactInfo(name="Jane Doe"),
        experience=[
            ExperienceItem(
                title="Developer",
                company="Acme",
                location="Kuala Lumpur",
                start_date="2020",
                end_date="Present",
                bullets=["Built things"],
            )
        ],
        projects=[
            ProjectItem(
                title="Migration Project",
                meta="Project within Acme role",
                start_date="Mar 2025",
                end_date="May 2025",
                bullets=["Migrated systems"],
            )
        ],
        education=[
            EducationItem(degree="BSc CS", institution="MIT", year="2015 - 2019")
        ],
    )


def test_docx_dates_share_heading_paragraph(tmp_path):
    path = tmp_path / "resume.docx"
    export_docx(_layout_resume(), str(path))
    paragraphs = [p.text for p in Document(path).paragraphs]

    assert any(
        "Developer" in text
        and "Acme" in text
        and "Kuala Lumpur" in text
        and "2020 \u2013 Present" in text
        and "\t" in text
        for text in paragraphs
    )
    assert any(
        "Migration Project" in text
        and "Mar 2025 \u2013 May 2025" in text
        and "\t" in text
        for text in paragraphs
    )
    assert any("Migration Project | Project within Acme role" in text for text in paragraphs)
    assert any(
        "BSc CS: MIT | 2015\u20132019" in text
        for text in paragraphs
    )
    assert "2020 \u2013 Present" not in paragraphs
    assert "Mar 2025 \u2013 May 2025" not in paragraphs


def test_pdf_dates_share_heading_baseline(tmp_path):
    path = tmp_path / "resume.pdf"
    export_pdf(_layout_resume(), str(path))

    with fitz.open(path) as pdf:
        words = [word for page in pdf for word in page.get_text("words")]

    def top_of(token: str) -> float:
        matches = [word for word in words if word[4] == token or word[4].startswith(token)]
        assert matches, f"Token not found in PDF: {token}"
        return matches[0][1]

    assert abs(top_of("Developer") - top_of("2020")) < 2
    assert abs(top_of("Migration") - top_of("Mar")) < 2
    all_text = " ".join(word[4] for word in words)
    assert "2015" in all_text and "2019" in all_text and "MIT" in all_text


def test_export_recovers_embedded_skills_section(tmp_path):
    resume = ResumeData(
        contact=ContactInfo(name="Jane Doe"),
        summary=(
            "Network engineer delivering secure multi-site projects. "
            "CORE TECHNICAL SKILLS "
            "Routing & Switching: BGP, OSPF, VLANs "
            "Security: Palo Alto, Fortinet "
            "SD-WAN: Huawei, Cato Networks"
        ),
        skills=[],
    )
    path = tmp_path / "recovered.docx"
    export_docx(resume, str(path))
    paragraphs = [p.text for p in Document(path).paragraphs]

    summary_text = paragraphs[paragraphs.index("PROFESSIONAL SUMMARY") + 1]
    assert summary_text == "Network engineer delivering secure multi-site projects."
    assert "CORE TECHNICAL SKILLS" not in summary_text
    assert "CORE TECHNICAL SKILLS" in paragraphs
    assert any(text.startswith("Routing & Switching: BGP") for text in paragraphs)
    assert any(text.startswith("Security: Palo Alto") for text in paragraphs)
    assert any(text.startswith("SD-WAN: Huawei") for text in paragraphs)


def test_exported_pdf_text_matches_resume_model(tmp_path):
    resume = _resume()
    path = tmp_path / "resume.pdf"
    export_pdf(resume, str(path))
    with fitz.open(path) as pdf:
        full_text = " ".join(
            word[4] for page in pdf for word in page.get_text("words")
        )
    full_lower = full_text.lower()
    assert "jane" in full_lower
    assert "doe" in full_lower
    assert "jane@example.com" in full_text
    assert "python" in full_lower
    assert "acme" in full_lower
    assert "developer" in full_lower
    assert "built" in full_lower
    assert "mit" in full_lower
    assert "aws" in full_lower


def test_summary_is_capped_at_300_words():
    resume = ResumeData(
        contact=ContactInfo(name="Jane Doe"),
        summary=" ".join(f"word{i}" for i in range(350)),
    )
    md = to_markdown(resume)
    summary = md.split("## Professional Summary", 1)[1].strip().splitlines()[0]
    assert len(summary.split()) <= 300


# ── Phase 8: ExportTheme + Templates ───────────────────────────────────


def test_export_theme_defaults():
    t = ExportTheme()
    assert t.name == "Classic"
    assert t.font_family == "Arial"
    assert t.accent_color == "185076"
    assert t.body_font_size == 9.0


def test_export_theme_rgb_properties():
    t = ExportTheme(accent_color="FF0000", secondary_color="00FF00", text_color="0000FF")
    assert t.accent_rgb == (0xFF, 0x00, 0x00)
    assert t.secondary_rgb == (0x00, 0xFF, 0x00)
    assert t.text_rgb == (0x00, 0x00, 0xFF)


def test_export_theme_float_properties():
    t = ExportTheme(accent_color="185076")
    r, g, b = t.accent_float()
    assert abs(r - 0x18 / 255) < 0.001
    assert abs(g - 0x50 / 255) < 0.001
    assert abs(b - 0x76 / 255) < 0.001


def test_templates_registry():
    assert "Classic" in TEMPLATES
    assert "Modern" in TEMPLATES
    assert "Compact" in TEMPLATES
    assert get_template("Classic") is CLASSIC_THEME
    assert get_template("Modern") is MODERN_THEME
    assert get_template("Compact") is COMPACT_THEME


def test_get_template_fallback():
    assert get_template("NonExistent") is CLASSIC_THEME
    assert get_template("") is CLASSIC_THEME


def test_modern_theme_differs_from_classic():
    assert MODERN_THEME.font_family == "Calibri"
    assert MODERN_THEME.accent_color != CLASSIC_THEME.accent_color
    assert MODERN_THEME.name == "Modern"


def test_compact_theme_has_smaller_margins():
    assert COMPACT_THEME.margin_top_bottom < CLASSIC_THEME.margin_top_bottom
    assert COMPACT_THEME.margin_left_right < CLASSIC_THEME.margin_left_right


def test_pdf_with_modern_theme(tmp_path):
    path = tmp_path / "modern.pdf"
    export_pdf(_resume(), str(path), theme=MODERN_THEME)
    with fitz.open(path) as pdf:
        full_text = " ".join(word[4] for page in pdf for word in page.get_text("words"))
    assert "JANE" in full_text


def test_pdf_with_compact_theme(tmp_path):
    path = tmp_path / "compact.pdf"
    export_pdf(_resume(), str(path), theme=COMPACT_THEME)
    with fitz.open(path) as pdf:
        assert len(pdf) == 1


def test_docx_with_modern_theme(tmp_path):
    path = tmp_path / "modern.docx"
    export_docx(_resume(), str(path), theme=MODERN_THEME)
    doc = Document(str(path))
    paragraphs = [p.text for p in doc.paragraphs]
    assert any("JANE" in p.upper() for p in paragraphs)


def test_docx_with_compact_theme(tmp_path):
    path = tmp_path / "compact.docx"
    export_docx(_resume(), str(path), theme=COMPACT_THEME)
    doc = Document(str(path))
    paragraphs = [p.text for p in doc.paragraphs]
    assert any("JANE" in p.upper() for p in paragraphs)


def test_pdf_target_pages_single_page(tmp_path):
    path = tmp_path / "single.pdf"
    export_pdf(_resume(), str(path), target_pages=1)
    with fitz.open(path) as pdf:
        assert len(pdf) <= 1


def test_pdf_target_pages_two(tmp_path):
    path = tmp_path / "two.pdf"
    export_pdf(_resume(), str(path), target_pages=2)
    with fitz.open(path) as pdf:
        assert len(pdf) <= 2


def test_export_backward_compatible_no_theme(tmp_path):
    """export_pdf/export_docx work without a theme argument (backward compat)."""
    pdf_path = tmp_path / "compat.pdf"
    export_pdf(_resume(), str(pdf_path))
    assert pdf_path.exists()

    docx_path = tmp_path / "compat.docx"
    export_docx(_resume(), str(docx_path))
    assert docx_path.exists()


def test_theme_font_family_applies_to_docx(tmp_path):
    path = tmp_path / "custom_font.docx"
    custom = ExportTheme(name="Test", font_family="Courier New")
    export_docx(_resume(), str(path), theme=custom)
    doc = Document(str(path))
    run = doc.paragraphs[0].runs[0]
    assert run.font.name == "Courier New"


def test_theme_font_family_applies_to_docx_heading(tmp_path):
    """The first paragraph is the name heading - check its font."""
    path = tmp_path / "heading_font.docx"
    custom = ExportTheme(name="Test", font_family="Consolas")
    export_docx(_resume(), str(path), theme=custom)
    doc = Document(str(path))
    first_run = doc.paragraphs[0].runs[0]
    assert first_run.font.name == "Consolas"


def test_all_templates_produce_valid_pdf(tmp_path):
    for name, theme in TEMPLATES.items():
        path = tmp_path / f"{name.lower()}.pdf"
        export_pdf(_resume(), str(path), theme=theme)
        with fitz.open(path) as pdf:
            assert len(pdf) >= 1
            text = " ".join(word[4] for page in pdf for word in page.get_text("words"))
            assert "JANE" in text


def test_all_templates_produce_valid_docx(tmp_path):
    for name, theme in TEMPLATES.items():
        path = tmp_path / f"{name.lower()}.docx"
        export_docx(_resume(), str(path), theme=theme)
        doc = Document(str(path))
        paragraphs = [p.text for p in doc.paragraphs]
        assert any("JANE" in p.upper() for p in paragraphs)
